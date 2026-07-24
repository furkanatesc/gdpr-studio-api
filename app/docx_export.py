"""Uretilen belgeler (Markdown) icin stilli .docx donusturucu."""

from __future__ import annotations

import io
import re

from docx import Document

from .docx_style import ACCENT, BASLIK_FONT, GOVDE_FONT, build_cover

_SEP = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def _is_sep(line: str) -> bool:
    return bool(_SEP.match(line))


def _parse_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _apply_default_font(doc) -> None:
    doc.styles["Normal"].font.name = GOVDE_FONT


def _heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = BASLIK_FONT
    run.font.color.rgb = ACCENT


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for c, htext in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.name = GOVDE_FONT
    for r, rowdata in enumerate(rows, start=1):
        for c, val in enumerate(rowdata):
            if c < len(headers):
                cell = table.rows[r].cells[c]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.name = GOVDE_FONT


def _render_body(doc, markdown_text: str) -> None:
    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        # Tablo: mevcut satir '|' iceriyor ve HEMEN sonraki satir ayrac.
        if "|" in stripped and i + 1 < n and _is_sep(lines[i + 1]):
            headers = _parse_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < n and lines[i].strip() and "|" in lines[i] and not _is_sep(lines[i]):
                rows.append(_parse_row(lines[i]))
                i += 1
            _add_table(doc, headers, rows)
            continue
        if stripped.startswith("#### "):
            _heading(doc, stripped[5:], 4)
        elif stripped.startswith("### "):
            _heading(doc, stripped[4:], 3)
        elif stripped.startswith("## "):
            _heading(doc, stripped[3:], 2)
        elif stripped.startswith("# "):
            _heading(doc, stripped[2:], 1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
        i += 1


def render_styled_docx(markdown_text: str, doc_type: str, cover_data: dict) -> bytes:
    """Kapak + stilli govde ureten stilli .docx; bayt dondurur."""
    doc = Document()
    _apply_default_font(doc)
    build_cover(
        doc,
        doc_type,
        veri_sorumlusu=cover_data.get("veri_sorumlusu"),
        ilgili_kisi=cover_data.get("ilgili_kisi"),
        site=cover_data.get("site"),
        tarih=cover_data.get("tarih"),
        versiyon=cover_data.get("versiyon"),
    )
    _render_body(doc, markdown_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_docx(markdown_text: str, title: str) -> bytes:
    """Kapaksiz stilli .docx (geriye donuk imza; kapak yerine baslik)."""
    doc = Document()
    _apply_default_font(doc)
    doc.add_heading(title, level=0)
    _render_body(doc, markdown_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

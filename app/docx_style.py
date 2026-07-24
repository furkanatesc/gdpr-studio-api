"""Stilli DOCX cikti icin palet, font sabitleri ve belge turune gore kapak."""

from __future__ import annotations

from docx.shared import Pt, RGBColor

from legal_core.prompt import ONAY_BEKLEYEN_PLACEHOLDER

ACCENT = RGBColor(0x3E, 0x7A, 0x72)
INK = RGBColor(0x1C, 0x1C, 0x1A)
MUTED = RGBColor(0x6B, 0x69, 0x63)
LINE = RGBColor(0xE4, 0xE2, 0xDB)

BASLIK_FONT = "Cormorant Garamond"
GOVDE_FONT = "Spectral"

_BASLIK = {
    "aydinlatma": "Kişisel Verilerin Korunması Kanunu Kapsamında Aydınlatma Metni",
    "cerez": "Çerez Politikası",
    "kayit": "Kişisel Veri İşleme Kaydı",
}


def _alan(doc, etiket: str, deger: str | None) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{etiket}: ")
    r1.font.name = GOVDE_FONT
    r1.font.color.rgb = MUTED
    r2 = p.add_run(deger if deger else ONAY_BEKLEYEN_PLACEHOLDER)
    r2.font.name = GOVDE_FONT
    r2.font.color.rgb = INK


def build_cover(
    doc,
    doc_type: str,
    *,
    veri_sorumlusu: str | None,
    ilgili_kisi: str | None,
    site: str | None,
    tarih: str | None,
    versiyon: str | None,
) -> None:
    """Belge turune gore kapak sayfasi kurar, sonuna page break ekler."""
    h = doc.add_paragraph()
    run = h.add_run(_BASLIK.get(doc_type, "Belge"))
    run.bold = True
    run.font.name = BASLIK_FONT
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    _alan(doc, "Veri Sorumlusu", veri_sorumlusu)
    if doc_type == "aydinlatma":
        _alan(doc, "İlgili Kişi", ilgili_kisi)
    if doc_type == "cerez":
        _alan(doc, "Site / Uygulama", site)
    _alan(doc, "Yürürlük Tarihi", tarih)
    _alan(doc, "Versiyon", versiyon)

    doc.add_page_break()

from docx import Document

from app.docx_style import build_cover


def _text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_kapak_aydinlatma_ilgili_kisi_var_site_yok():
    doc = Document()
    build_cover(doc, "aydinlatma", veri_sorumlusu="ACME A.S.", ilgili_kisi="Calisan",
                site=None, tarih="24.07.2026", versiyon="1.0")
    t = _text(doc)
    assert "ACME A.S." in t
    assert "İlgili Kişi" in t and "Calisan" in t
    assert "Site" not in t
    assert "24.07.2026" in t
    assert "1.0" in t


def test_kapak_cerez_site_var_ilgili_kisi_yok():
    doc = Document()
    build_cover(doc, "cerez", veri_sorumlusu="ACME", ilgili_kisi=None,
                site="ornek.com", tarih="24.07.2026", versiyon="1.0")
    t = _text(doc)
    assert "Site" in t and "ornek.com" in t
    assert "İlgili Kişi" not in t


def test_kapak_kayit_ilgili_kisi_ve_site_yok():
    doc = Document()
    build_cover(doc, "kayit", veri_sorumlusu="ACME", ilgili_kisi=None,
                site=None, tarih="24.07.2026", versiyon="1.0")
    t = _text(doc)
    assert "ACME" in t
    assert "İlgili Kişi" not in t
    assert "Site" not in t
    assert "Yürürlük Tarihi" in t


def test_kapak_bos_veri_sorumlusu_placeholder():
    doc = Document()
    build_cover(doc, "aydinlatma", veri_sorumlusu=None, ilgili_kisi=None,
                site=None, tarih="24.07.2026", versiyon="1.0")
    assert "[Avukat tarafından doldurulacak]" in _text(doc)

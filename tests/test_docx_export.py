import io
import io as _io
import zipfile

import docx

from app.docx_export import render_docx, render_styled_docx

MARKDOWN = (
    "## Baslik\n\nBir paragraf.\n\n- madde bir\n- madde iki\n\n"
    "Bu cikti avukat incelemesine tabidir, hukuki tavsiye yerine gecmez."
)
TITLE = "Aydinlatma Metni"


def test_render_docx_returns_nonempty_bytes():
    result = render_docx(MARKDOWN, TITLE)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_render_docx_is_valid_zip_with_document_xml():
    result = render_docx(MARKDOWN, TITLE)
    with zipfile.ZipFile(io.BytesIO(result)) as zf:
        assert "word/document.xml" in zf.namelist()


def test_render_docx_contains_title_content_and_disclaimer():
    result = render_docx(MARKDOWN, TITLE)
    doc = docx.Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert TITLE in full_text
    assert "madde bir" in full_text
    assert "avukat incelemesine tabidir" in full_text


TABLO_MD = "| Bilgi | Değer |\n| --- | --- |\n| Ad | ACME |\n| Vergi No | 123 |"


def test_markdown_tablo_word_tablosuna_donusur():
    b = render_styled_docx(
        TABLO_MD, "kayit",
        {"veri_sorumlusu": "ACME", "tarih": "24.07.2026", "versiyon": "1.0"},
    )
    doc = docx.Document(_io.BytesIO(b))
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert len(tbl.columns) == 2
    assert len(tbl.rows) == 3  # baslik + 2 satir
    assert tbl.rows[0].cells[0].text == "Bilgi"
    assert tbl.rows[1].cells[1].text == "ACME"


def test_render_styled_docx_kapak_ve_govde_birlikte():
    b = render_styled_docx(
        "## 1. Tanımlar\n\nBir metin.", "aydinlatma",
        {"veri_sorumlusu": "ACME", "ilgili_kisi": "Calisan",
         "tarih": "24.07.2026", "versiyon": "1.0"},
    )
    doc = docx.Document(_io.BytesIO(b))
    t = "\n".join(p.text for p in doc.paragraphs)
    assert "ACME" in t and "Calisan" in t  # kapak
    assert "Tanımlar" in t                  # govde basligi


def test_render_docx_hala_calisir_ve_tablo_isler():
    # Mevcut imza korunur; tablo destegi de kazanir.
    b = render_docx(TABLO_MD, "Baslik")
    doc = docx.Document(_io.BytesIO(b))
    assert len(doc.tables) == 1
